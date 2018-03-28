from django.shortcuts import render, render_to_response
from django import forms
from django.http import HttpResponse
from .models import NormalUser
import docx
import codecs
from win32com import client
from pdfminer.pdfparser import PDFParser,PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal,LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
import pythoncom
from . import models

#路径统一在这里修改
store_path = r'F:\github\WBFL\uploadpath\\'
file_path =r'F:\github\WBFL\TianWen\media\upload\\'
'''
store_path = 'F:\\github\\WBFL\\uploadpath\\'
file_path = 'F:\\github\\WBFL\\uploadpath\\'
'''

# Create your views here.
class NormalUserForm(forms.Form):
# form的定义和model类的定义很像
    username=forms.CharField()
    headImg=forms.FileField()


#在View中使用已定义的Form方法
def registerNormalUser(request):
    #刚显示时调用GET方法
    if request.method=="POST":
        uf = NormalUserForm(request.POST,request.FILES)  #刚显示时，实例化表单（是否有数据）
        if uf.is_valid():  #验证数据是否合法，当合法时可以使用cleaned_data属性。
            #用来得到经过'clean'格式化的数据，将所提交过来的数据转化成合适的Python的类型。
            username = uf.cleaned_data['username']
            headImg = uf.cleaned_data['headImg']
            #write in database
            normalUser=NormalUser()  #实例化NormalUser对象
            normalUser.username = username
            normalUser.headImg = headImg
            normalUser.save()#保存到数据库表中
            suffix = headImg.name[headImg.name.find('.'):len(headImg.name) + 1]
            path= file_path + headImg.name
            if suffix == '.docx':
                read_docx(request,path)
            elif suffix == '.pdf':
                read_pdf(request,path)
            elif suffix == '.ppt':
                read_ppt(request, path)
            elif suffix == '.doc':
                read_doc(request, path)
            return HttpResponse('Upload Succeed!')#重定向显示内容（跳转后内容）
    else:
        uf=NormalUserForm()#刚显示时，实例化空表单
    return render(request,'Upload/register.html',{'uf':uf})#只有刚显示时才起作用


def read_docx(request, path):  # 阅读docx文档,path为文件的路径
    doc = docx.Document(path)  # 根据路径打开一个docx对象
    file_name = path[path.rfind('\\') + 1:path.find('.')]  # 提取路径中的文件名,不包含后缀
    outfile = open(store_path + file_name + '_temp.txt', 'w',encoding = 'utf-8')  # 创建一个临时的txt文件
    for i in range(len(doc.paragraphs)):  # 将docx文档里的数据写入txt文档中
        outfile.write(doc.paragraphs[i].text + '\n')
    # 读取docx文档中的表格中的字符，加在正文后面
    for i in range(len(doc.tables)):  # 获取所有表格的数目
        for j in range(len(doc.tables[i].rows)):  # 获取单个表格的行数
            for k in range(len(doc.tables[i].columns)):  # 获取单个表格的列数
                outfile.write(doc.tables[i].cell(j, k).text)  # 将单元格的内容写入txt文本中
    outfile.close()  # 结束写入


def read_pdf(request, path):
    fp = open(path, 'rb')  # 以二进制读模式打开
    file_name = path[path.rfind('\\') + 1:path.find('.')]  # 提取路径中的文件名,不包含后缀
    outfile = open(store_path + file_name + '_temp.txt', 'w',encoding = 'utf-8')  # 创建一个临时的txt文件
    # 用文件对象来创建一个pdf文档分析器
    praser = PDFParser(fp)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 连接分析器 与文档对象
    praser.set_document(doc)
    doc.set_parser(praser)

    # 提供初始化密码
    # 如果没有密码 就创建一个空的字符串
    doc.initialize()

    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        # 创建PDf 资源管理器 来管理共享资源
        rsrcmgr = PDFResourceManager()
        # 创建一个PDF设备对象
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解释器对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        # 循环遍历列表，每次处理一个page的内容
        for page in doc.get_pages():  # doc.get_pages() 获取page列表
            interpreter.process_page(page)
            # 接受该页面的LTPage对象
            layout = device.get_result()
            # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等 想要获取文本就获得对象的text属性，
            for x in layout:
                if (isinstance(x, LTTextBoxHorizontal)):
                    results = x.get_text()
                    print(results)
                    outfile.write(results + '\n')
    outfile.close()


def read_ppt(request, path):
    pythoncom.CoInitialize()
    file_name = path[path.rfind('\\') + 1:path.find('.')]  # 提取路径中的文件名,不包含后缀
    ppt = client.Dispatch('PowerPoint.Application')
    ppt.Visible = 1
    pptSel = ppt.Presentations.Open(path)
    client.gencache.EnsureDispatch('PowerPoint.Application')
    f = open(store_path + file_name + '_temp.txt', 'w',encoding = 'utf-8')   # 创建一个临时的txt文件
    slide_count = pptSel.Slides.Count
    for i in range(1, slide_count + 1):
        shape_count = pptSel.Slides(i).Shapes.Count
        for j in range(1, shape_count + 1):
            if pptSel.Slides(i).Shapes(j).HasTextFrame:
                s = pptSel.Slides(i).Shapes(j).TextFrame.TextRange.Text
                f.write(s.encode('utf-8').decode() + "\n")
    f.close()
    ppt.Quit()


def read_doc(request,path):  #将doc文档转换成docx文档并读取,path为文件的路径
        pythoncom.CoInitialize()
        file_name = path[path.rfind('\\') + 1:path.find('.')]  #提取路径中的文件名,不包含后缀
        word = client.Dispatch("word.Application")
        doc = word.Documents.Open(path)
        doc.SaveAs(file_path + file_name + '.docx', 16)
        doc.Close()
        word.Quit()
        read_docx(request,file_path + file_name + '.docx')



